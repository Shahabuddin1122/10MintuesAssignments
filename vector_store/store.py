import os
from typing import List
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.document_loaders import TextLoader, PyPDFLoader
from langchain.vectorstores import FAISS
from vector_store.embedder import get_embedder


# Loaders
def load_txt(path: str) -> List[Document]:
    return TextLoader(path, encoding="utf-8").load()


def load_pdf(path: str) -> List[Document]:
    return PyPDFLoader(path).load()


def load_docx(path: str) -> List[Document]:
    full_text = "\n".join([para.text.strip() for para in DocxDocument(path).paragraphs if para.text.strip()])
    return [Document(page_content=full_text, metadata={"source": path})]


def load_documents(paths: List[str]) -> List[Document]:
    all_docs = []
    for path in paths:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            all_docs.extend(load_txt(path))
        elif ext == ".pdf":
            all_docs.extend(load_pdf(path))
        elif ext == ".docx":
            all_docs.extend(load_docx(path))
    return all_docs


def save_vectorstore(docs: List[Document], path: str):
    embedder = get_embedder()
    vectordb = FAISS.from_documents(docs, embedder)
    vectordb.save_local(path)
    print(f"Vectorstore saved to: {path}")
